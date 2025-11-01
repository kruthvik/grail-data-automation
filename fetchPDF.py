from docx import Document
from docx2pdf import convert
from random import randint
import requests as req
from bs4 import BeautifulSoup as bs
import random 
import re
import os
import glob

import os
import glob
import json as js
from datetime import datetime
from findAgency import abbreviateAgency

formatCommenter = lambda commenter: re.sub(r"(?i)^(comment submitted by|comment from)\s*", "", commenter).strip()

<<<<<<< HEAD
def removeItems(folder="comments/", removeMetadata=True):
    files = glob.glob(os.path.join(folder, "*"))
    for f in files:
        if os.path.isfile(f):
            try:
                os.remove(f)
                print(f"Deleted: {f}")
            except Exception as e:
                print(f"Failed to delete {f}: {e}")
    print("All files removed from", folder)
    
    if removeMetadata:
        with open("./metadata.json", "w", encoding="utf-8") as f:
            js.dump({}, f)
            print("Metadata reset.")
        
def loadJSON():
    try:
        with open("./metadata.json", "r", encoding="utf-8") as f:
=======
def setupFolders(folder=""):
    comments = "comments/" + folder
    os.makedirs(comments, exist_ok=True)

    logs = "logs/" + folder
    os.makedirs(logs, exist_ok=True)
    
    os.removeItems(folder=comments, removeMetadata=False)
    os.removeItems(folder=logs, removeMetadata=False)

    with open(f"./comments/{folder}/metadata.json", "w", encoding="utf-8") as f:
        js.dump({}, f)

def loadJSON(folder=""):
    try:
        with open(f"./comments/{folder}/metadata.json", "r", encoding="utf-8") as f:
>>>>>>> 3b29589 (updated)
            data = js.load(f)
            return data
    except FileNotFoundError:
        return {}
<<<<<<< HEAD

def make_filename(notice_shortname, commenter, posted_date, agency_full):
=======
    
def writeJSON(data, folder=""):
    with open(f"./comments/{folder}/metadata.json", "w", encoding="utf-8") as f:
        js.dump(data, f)

def make_filename(commenter, posted_date, agency_full):
>>>>>>> 3b29589 (updated)
    agency_full = agency_full.lower()
    short = abbreviateAgency(agency_full)

    # Clean commenter
    commenter_clean = re.sub(r"[^\w]", "", formatCommenter(commenter))

    # Parse date to YYYYMMDD
    try:
        dt = datetime.strptime(posted_date, "%b %d, %Y")
<<<<<<< HEAD
        date_str = dt.strftime("%Y%m%d")
    except Exception:
        try:
            dt = datetime.fromisoformat(posted_date)
            date_str = dt.strftime("%Y%m%d")
=======
        date_str = dt.strftime("%m/%d/%Y")
    except Exception:
        try:
            dt = datetime.fromisoformat(posted_date)
            date_str = dt.strftime("%m/%d/%Y")
>>>>>>> 3b29589 (updated)
        except Exception:
            date_str = "unknownDate"

    return f"{short}_{commenter_clean}_{date_str}"

<<<<<<< HEAD
def createPDF(info, date, commenter, agency, comment_id):
    commenter = formatCommenter(commenter)
    filename_base = make_filename(comment_id, commenter, date, agency)
=======
def createPDF(info, date, commenter, agency, comment_id, docNum):
    commenter = formatCommenter(commenter)
    filename_base = make_filename(commenter, date, agency)
>>>>>>> 3b29589 (updated)

    doc = Document()

    # Add a title
    doc.add_heading(f"Comment from {commenter}", level=0)

    # Add a paragraph
    doc.add_paragraph(f"Date: {date}\nAgency: {agency}\nID: {comment_id}")

    doc.add_paragraph(info)

    # Save the document
    doc.save(f"./comments/{filename_base}.docx")

    print(f"DOCX file saved as {filename_base}.docx")

<<<<<<< HEAD
    convert(f"./comments/{filename_base}.docx", f"./comments/{filename_base}.pdf")
    
    json = loadJSON()
=======
    convert(f"./comments/{docNum}/{filename_base}.docx", f"./comments/{docNum}/{filename_base}.pdf")

    json = loadJSON(folder=docNum)
>>>>>>> 3b29589 (updated)
    json[filename_base] = {
        "date": date,
        "commenter": commenter,
        "agency": agency,
        "comment_id": comment_id
    }
<<<<<<< HEAD
    
    with open(f"./metadata.json", "w", encoding="utf-8") as f:
        js.dump(json, f)
        print("updated")
=======

    writeJSON(json, folder=docNum)
    print(f"Updated metadata.json for {docNum}")
>>>>>>> 3b29589 (updated)
    
    print(f"PDF file saved as {filename_base}.pdf")
    
    os.remove(f"./comments/{filename_base}.docx")

<<<<<<< HEAD
def downloadPDF(url, date, commenter, agency, comment_id):
=======
def downloadPDF(url, date, commenter, agency, comment_id, docNum):
>>>>>>> 3b29589 (updated)
    commenter = formatCommenter(commenter)
    
    r = req.get(url)
    fname = make_filename(comment_id, commenter, date, agency)
<<<<<<< HEAD
    
    with open(f"./comments/{fname}.pdf", "wb") as f:
        f.write(r.content)

    json = loadJSON()
=======

    with open(f"./comments/{docNum}/{fname}.pdf", "wb") as f:
        f.write(r.content)

    json = loadJSON(folder=docNum)
>>>>>>> 3b29589 (updated)
    json[fname] = {
        "date": date,
        "commenter": commenter,
        "agency": agency,
        "comment_id": comment_id
    }
<<<<<<< HEAD
    
    with open(f"./metadata.json", "w", encoding="utf-8") as f:
        js.dump(json, f)
        print("updated")

    print(f"PDF file saved as {make_filename(comment_id, commenter, date, agency)}.pdf")
=======

    writeJSON(json, folder=docNum)
    print(f"Updated metadata.json for {docNum}")

    print(f"PDF file saved as {make_filename(commenter, date, agency)}.pdf")
>>>>>>> 3b29589 (updated)
