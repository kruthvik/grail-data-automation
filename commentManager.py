from docx import Document
from docx2pdf import convert
from random import randint
import requests as req
from bs4 import BeautifulSoup as bs
import random 
import re

import os
import shutil
import json as js
from datetime import datetime
from NoticeAnalyzer import NoticeAnalyzer

import fitz

formatCommenter = lambda commenter: re.sub(r"(?i)^(comment submitted by|comment from)\s*", "", commenter).strip()

class CommentManager:
    def __init__(self, logger=None, documentID=None):
        self.logger = logger
        self.documentID = documentID
        self.folderPath = f"./comments/{self.documentID}/" if documentID else None

    @staticmethod
    def removeItems(folder="", logger=None):
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:    
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                if logger:
                    logger.log(f'Failed to delete {file_path}. Reason: {e}', level="WARNING")
                else:
                    print(f'ERROR: Failed to delete {file_path}. Reason: {e}')
      
    @staticmethod
    def make_filename(commenter, posted_date, agency_full, commenterId, logger=None):
        agency_full = agency_full.lower()
        short = NoticeAnalyzer.abbreviateAgency(agency_full)

        # Clean commenter
        commenter_clean = re.sub(r"[^\w]", "", formatCommenter(commenter))

        # Parse date to YYYYMMDD
        try:
            dt = datetime.strptime(posted_date, "%b %d, %Y")
            date_str = dt.strftime("%Y%m%d")
        except Exception:
            try:
                dt = datetime.fromisoformat(posted_date)
                date_str = dt.strftime("%Y%m%d")
            except Exception:
                if logger:
                    logger.log(f"Failed to parse date {posted_date}. Using 'unknownDate'.", level="WARNING")
                else:
                    print(f"ERROR: Failed to parse date {posted_date}. Using 'unknownDate'.")
                date_str = "unknownDate"

        return f"{short}__{commenter_clean}__{date_str}__{commenterId}"              

    def setupFolders(self):
        os.makedirs(self.folderPath, exist_ok=True)

        self.logger.log(f"Setup folders for {self.folderPath}", level="INFO")

        CommentManager.removeItems(folder=self.folderPath, logger=self.logger)

        with open(f"{self.folderPath}/metadata.json", "w", encoding="utf-8") as f:
            js.dump({}, f)
            
    def loadJSON(self):
        try:
            with open(f"{self.folderPath}/metadata.json", "r", encoding="utf-8") as f:
                data = js.load(f)
                return data
        except FileNotFoundError:
            if self.logger:
                self.logger.log(f"metadata.json not found in folder {self.folderPath}. Returning empty dict.", level="WARNING")
            else:
                print(f"ERROR: metadata.json not found in folder {self.folderPath}. Returning empty dict.")
            return {}

    def writeJSON(self, data):
        with open(f"{self.folderPath}/metadata.json", "w", encoding="utf-8") as f:
            js.dump(data, f)

    def createPDF(self, info, date, commenter, agency, comment_id, numAttachments=0, attachmentNum=0):
        commenter = formatCommenter(commenter)
        filename_base = CommentManager.make_filename(commenter, date, agency, comment_id, logger=self.logger)

        doc = Document()

        # Add a title
        doc.add_heading(f"Comment from {commenter}", level=0)

        # Add a paragraph
        doc.add_paragraph(f"Date: {date}\nAgency: {agency}\nID: {comment_id}")

        doc.add_paragraph(info)

        if numAttachments > 1:
            filename_base = f"{filename_base}_attachment{attachmentNum}"

        # Save the document
        doc.save(f"{self.folderPath}/{filename_base}.docx")

        self.logger.log(f"DOCX file saved as {filename_base}.docx")

        convert(f"{self.folderPath}/{filename_base}.docx", f"{self.folderPath}/{filename_base}.pdf")

        json = self.loadJSON()
        
        json[filename_base] = {
            "date": date,
            "commenter": commenter,
            "agency": agency,
            "comment_id": comment_id
        }

        self.writeJSON(json)
        
        self.logger.log(f"Updated metadata.json for {comment_id}")
        self.logger.log(f"PDF file saved as {filename_base}.pdf")

        os.remove(f"{self.folderPath}/{filename_base}.docx")

    def flattenPDF(self, input_path: str):
        """
        Extracts the main PDF from a PDF Portfolio (if applicable) and overwrites
        the original file with the extracted one. If the file isn't a portfolio,
        it is left unchanged.
        """
        try:
            doc = fitz.open(input_path)
        except Exception as e:
            print(f"❌ Failed to open {input_path}: {e}")
            return

        emb_count = doc.embfile_count()
        print(f"Found {emb_count} embedded files in {input_path}")

        # If no embedded files, it's already a normal PDF
        if emb_count == 0:
            doc.close()
            print("✅ No embedded files — already a regular PDF.")
            return

        extracted_main = None

        for i in range(emb_count):
            try:
                try:
                    info = doc.embfile_info(i)
                except UnicodeEncodeError:
                    info = {}
                
                data = doc.embfile_get(i)

                # Detect PDFs
                if data.startswith(b"%PDF"):
                    extracted_main = data
                    break  # usually the first embedded PDF is the main document
            except Exception as e:
                print(f"⚠️ Error extracting embedded file #{i}: {e}")

        doc.close()

        if extracted_main:
            with open(input_path, "wb") as f:
                f.write(extracted_main)
            print(f"✅ Flattened portfolio and replaced original: {input_path}")
        else:
            print("⚠️ No embedded PDF found in portfolio.")

    def downloadPDF(self, url, date, commenter, agency, comment_id, numAttachments=0, attachmentNum=0):
        
        commenter = formatCommenter(commenter)
        
        r = req.get(url)
        
        if r.status_code != 200:
            self.logger.log(f"Failed to download {url} ({r.status_code})", level="ERROR")
            return

        fname = CommentManager.make_filename(commenter, date, agency, comment_id, logger=self.logger)
        suffix = url.split('.')[-1].lower()
        
        if numAttachments > 1:
            folderPrefix = f"{self.folderPath}/{fname}/"

            fname = f"{fname}_attachment{attachmentNum}"

        with open(f"{self.folderPath}/{fname}.{suffix}", "wb") as f:
            f.write(r.content)

        json = self.loadJSON()
        json[fname] = {
            "date": date,
            "commenter": commenter,
            "agency": agency,
            "comment_id": comment_id
        }

        self.writeJSON(json)
        self.logger.log(f"Updated metadata.json for {comment_id}")

        self.logger.log(f"PDF file saved as {fname}.{suffix}")
        
        if suffix == "pdf":
            self.flattenPDF(f"{self.folderPath}/{fname}.pdf")
